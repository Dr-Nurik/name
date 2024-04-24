from .models import (SpecializationCategory, Doctor,
                     Service, Reception,
                     MedicalEquipment, Diagnosis,
                     UserProfile, Masseur, TrainingEquipment, Trainer)
from .forms import (ReceptionForm, ServiceForm,
                    TimePeriodForm, DiagnosisForm,
                    ReceptionEditForm)
from django.views.generic.edit import FormView
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.contrib.auth.mixins import UserPassesTestMixin
from django.contrib.admin.views.decorators import staff_member_required
from django.utils.decorators import method_decorator
from django.shortcuts import render, redirect, get_object_or_404
from datetime import date, timedelta
from django.db.models import Count, F
from collections import defaultdict
from django.utils import timezone  # Импортируйте timezone из django.utils
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.http import HttpResponse
from datetime import datetime
from urllib.parse import quote
from django.urls import reverse_lazy
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .serializers import ReceptionSerializer, ServiceSerializer, DoctorSerializer, MasseurSerializer, TrainingEquipmentSerializer
from datetime import datetime, timedelta
from .forms import DateRangeForm
from django.db.models import Q
from django.utils.translation import gettext as _
from django.views.generic.edit import FormView
import json

import logging
logger = logging.getLogger(__name__)

class AdminRequiredMixin(UserPassesTestMixin):
    def test_func(self):
        return self.request.user.is_staff

def reception_confirmation_view(request):
    # Здесь ваша логика представления
    return render(request, 'main/reception_confirmation.html')
from django.shortcuts import redirect
from django.views.generic.edit import FormView
from .forms import ReceptionForm
from .models import Reception, Service, Doctor, Masseur, TrainingEquipment

class ReceptionView(FormView):
    form_class = ReceptionForm
    template_name = 'main/main.html'

    def form_valid(self, form):
        fcd = form.cleaned_data
        logger.info(f"Received form data: {fcd}")

        new_reception = Reception(
            date=fcd.get('date'),
            fullname_patient=fcd.get('fullname_patient'),
            patient_contact=fcd.get('patient_contact'),
            doctor=fcd.get('doctor'),
            masseur=fcd.get('masseur'),
            trainer=fcd.get('trainer')
        )
        new_reception.save()
        new_reception.services.set(fcd['services'])
        logger.info(f"Reception created successfully: {new_reception.id}")
        return redirect('main:reception_confirmation')

    def form_invalid(self, form):
        logger.error(f"Form is invalid: {form.errors}")
        return super().form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['doctors_list'] = Doctor.objects.all()
        context['masseur_list'] = Masseur.objects.all()
        context['trainer_list'] = TrainingEquipment.objects.all()
        context['trainers_list'] = Trainer.objects.all()
        data_types_json = json.dumps([str(_("Здоровый След")), str(_("Ортопедический центр!"))])
        context['data_types_json'] = data_types_json

        return context


class DoctorDetailView(DetailView):
    model = Doctor
    template_name = 'main/treatment/doctor_detail.html'  # Укажите путь к вашему шаблону

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['doctors_list'] = Doctor.objects.all()  # Добавление списка докторов в контекст
        return context

class ServiceUpdateView(UpdateView):
    model = Service
    form_class = ServiceForm
    template_name = 'main/service_update_form.html'
    success_url = reverse_lazy('service_list')

from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

from django.db.models import Q
from .models import Reception, Diagnosis  # Import your models here

def confirmed_and_arrived_receptions(request):
    search_query = request.GET.get('search', '')
    receptions_query = Reception.objects.filter(comed=True, is_confirmed=True)

    if search_query:
        receptions_query = receptions_query.filter(
            Q(fullname_patient__icontains=search_query) |
            Q(diagnosis__diagnosis_text__icontains=search_query) |
            Q(diagnosis__treatment_text__icontains=search_query)
        ).distinct()

    # Set up pagination
    paginator = Paginator(receptions_query, 20)  # 20 receptions per page
    page = request.GET.get('page')

    try:
        receptions = paginator.page(page)
    except PageNotAnInteger:
        # If page is not an integer, deliver first page.
        receptions = paginator.page(1)
    except EmptyPage:
        # If page is out of range (e.g., 9999), deliver last page of results.
        receptions = paginator.page(paginator.num_pages)

    diagnosis_list = Diagnosis.objects.filter(reception__in=receptions).select_related('reception')

    context = {
        'receptions': receptions,
        'diagnosis_list': diagnosis_list,
        'search_query': search_query,
    }
    return render(request, 'main/receptions/conf_arr_receptions_list.html', context)


from collections import defaultdict
from django.shortcuts import render
from django.utils import timezone
from .models import Reception

def confirmed_receptions_for_trainer(request):
    receptions_query = Reception.objects.filter(
        is_confirmed=True,
        trainer__isnull=True,
        comed=False,
        date__gte=timezone.now().date()
    ).order_by('date', 'time')

    receptions_by_date = defaultdict(list)
    for reception in receptions_query:
        receptions_by_date[reception.date].append(reception)

    context = {
        'receptions_by_date': dict(receptions_by_date),
    }
    return render(request, 'main/receptions/trainer_receptions_list.html', context)


class ServiceListView(ListView):
    model = Service  # Указываем модель, данные которой будут отображаться
    template_name = 'main/service_list.html'  # Указываем шаблон для отображения

class ServiceDetailView(DetailView):
    model = Service
    template_name = 'main/service_detail.html'  # Template to view details of a service

class ServiceCreateView(CreateView):
    model = Service
    form_class = ServiceForm
    template_name = 'main/service_form.html'  # Template for creating a new service
    success_url = reverse_lazy('service_list')  # Redirect URL after successful creation

def diagnosis_and_treatment_plan(request, reception_id):
    reception = get_object_or_404(Reception, pk=reception_id)
    diagnosis = Diagnosis.objects.filter(reception=reception).first()

    if request.method == 'POST':
        diagnosis_form = DiagnosisForm(request.POST, instance=diagnosis)

        if diagnosis_form.is_valid():
            saved_diagnosis = diagnosis_form.save(commit=False)
            saved_diagnosis.reception = reception
            saved_diagnosis.save()

            # Check if diagnosis and treatment plan are set
            if saved_diagnosis.diagnosis_text and saved_diagnosis.treatment_text:
                reception.comed = True
                reception.save()

            return redirect('main:reception_list')  # Redirect after successful save
    else:
        diagnosis_form = DiagnosisForm(instance=diagnosis)

    return render(request, 'main/treatment/diagnosis_and_treatment_plan.html', {
        'diagnosis_form': diagnosis_form,
        'reception': reception
    })

class ReceptionListView(ListView):
    model = Reception
    template_name = 'reception_list.html'
    context_object_name = 'receptions'

    def get_queryset(self):
        # Фильтруем подтвержденные записи с датой больше или равной текущей
        return Reception.objects.filter(is_confirmed=True, comed=False, date__gte=timezone.now().date()).order_by('date', 'time')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        receptions_by_date = defaultdict(list)

        for reception in self.get_queryset():
            receptions_by_date[reception.date].append(reception)

        context['receptions_by_date'] = dict(receptions_by_date)
        return context

def export_to_excel(request):
    # Получаем только подтвержденные записи
    receptions = Reception.objects.filter(is_confirmed=True, date__gte=datetime.today())

    # Создаем список для данных DataFrame
    data = []
    for reception in receptions:
        reception_data = {
            'Дата': reception.date,
            'Время': reception.time,
            'Пациент': reception.fullname_patient,
            'Услуги': ', '.join([service.name for service in reception.services.all()]),
            'Доктор': reception.doctor.doctors_name.full_name,
            'Тренажер': reception.trainer.name if reception.trainer else '',
            'Массажист': reception.masseur.masseur_name.full_name if reception.masseur else '',
            # Другие поля по необходимости
        }
        data.append(reception_data)

    df = pd.DataFrame(data)
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="receptions.xls"'

    with pd.ExcelWriter(response) as writer:
        df.to_excel(writer, index=False)

    return response

def export_to_word_by_date(request, date):
    receptions = Reception.objects.filter(date=date, is_confirmed=True).select_related('doctor', 'masseur', 'trainer').prefetch_related('services')

    # Создание документа Word
    doc = Document()
    # Добавление шапки с логотипом
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()

    run.add_picture('staticfiles/img/it_logo.png', width=Inches(2))  # Укажите путь к файлу логотипа

    # Заголовок документа
    title = doc.add_paragraph('Записи на ' + str(date))
    title.runs[0].font.name = 'Times New Roman'
    title.runs[0].font.size = Pt(14)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Добавление таблицы
    table = doc.add_table(rows=1, cols=4)  # Таблица с одной строкой заголовков и пятью колонками
    table.style = 'Table Grid'

    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Время'
    hdr_cells[1].text = 'Пациент'
    hdr_cells[2].text = 'Врач'
    hdr_cells[3].text = 'Услуга'

    # Наполнение таблицы данными
    for reception in receptions:
        row_cells = table.add_row().cells
        row_cells[0].text = str(reception.time)

        # Получение имени пациента
        fullname_patient = reception.fullname_patient
        row_cells[1].text = fullname_patient

        # Получение имени врача
        doctor_name = reception.doctor.doctors_name.full_name if reception.doctor else 'Не указано'
        row_cells[2].text = doctor_name

        # Получение списка услуг
        row_cells[3].text = ', '.join([service.name for service in reception.services.all()])

        # Настройка стилей ячеек
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    # Добавление контактных данных в футер
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = 'Ваши контактные данные здесь'  # Замените на ваши контактные данные
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    filename = f"Записи за {date}.docx"
    quoted_filename = quote(filename)

    # Сохранение документа в HttpResponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{quoted_filename}'

    doc.save(response)

    return response
from django.shortcuts import render
from datetime import datetime, timedelta
from .forms import DateRangeForm  # Убедитесь, что импортировали форму правильно
from django.db.models import Count, F

@staff_member_required
# def annual_report(request):
#     current_year = datetime.now().year
#     receptions = Reception.objects.filter(date__year=current_year, comed=True, is_confirmed=True)
#
#     # Получаем уникальных пациентов, которые посетили врача и подтвердили приход
#     patients_visited_doctor = Reception.objects.filter(
#         date__year=current_year,
#         doctor__isnull=False,
#         comed=True,
#         is_confirmed=True
#     ).values_list('patient_name_id', flat=True).distinct()
#
#     # Считаем, сколько из этих пациентов также посетили тренажерный зал
#     gym_attendance_count = Reception.objects.filter(
#         patient_name_id__in=patients_visited_doctor,
#         trainer__isnull=False,
#         comed=True,
#         is_confirmed=True,
#         date__gte=F('patient_name__reception__date')
#     ).distinct().count()
#
#
#     # Статистика по услугам
#     service_stats = receptions.values('services__name').annotate(count=Count('services')).order_by('-count')
#     service_data = {
#         'labels': [item['services__name'] for item in service_stats] if service_stats else [],
#         'data': [item['count'] for item in service_stats] if service_stats else [],
#     }
#
#     context = {
#         'current_year': current_year,
#         'patient_count': receptions.count(),
#         'gym_attendance_count': gym_attendance_count,
#         'service_stats': service_stats,
#         'service_data': service_data,
#     }
#
#     return render(request, 'main/treatment/annual_report.html', context)

def annual_report(request):
    today = datetime.now()
    form = DateRangeForm(request.GET)
    start_date, end_date = None, None

    if form.is_valid():
        period = form.cleaned_data.get('period')
        if period == 'year':
            start_date = datetime(today.year, 1, 1)
            end_date = today
        elif period == 'half_year':
            start_date = today - timedelta(days=182)
            end_date = today
        elif period == 'quarter':
            start_date = today - timedelta(days=90)
            end_date = today
        elif period == 'month':
            start_date = today - timedelta(days=30)
            end_date = today
        elif period == 'week':
            start_date = today - timedelta(days=7)
            end_date = today
        elif period == 'custom':
            start_date = form.cleaned_data.get('start_date')
            end_date = form.cleaned_data.get('end_date')

            # Фильтрация приемов
    receptions = Reception.objects.filter(
        date__range=(start_date, end_date),
        comed=True, is_confirmed=True
    ) if start_date and end_date else Reception.objects.filter(
        date__year=today.year,
        comed=True, is_confirmed=True
    )

    # Подсчет общего количества записей
    total_patients = receptions.count()

    print('Total appointments:', total_patients)

    # Статистика по услугам
    service_stats = receptions.values('services__name').annotate(count=Count('services')).order_by('-count')
    service_data = {
        'labels': [item['services__name'] for item in service_stats] if service_stats else [],
        'data': [item['count'] for item in service_stats] if service_stats else [],
    }

    # Статистика посещения тренажерного зала
    gym_attendance_count = receptions.filter(trainer__isnull=False).count()

    context = {
        'form': form,
        'current_year': today.year,
        'total_patients_count': total_patients,
        'gym_attendance_count': gym_attendance_count,
        'service_stats': service_stats,
        'service_data': service_data,
        'start_date': start_date,
        'end_date': end_date,
    }
    return render(request, 'main/treatment/annual_report.html', context)

from django.http import HttpResponse
from docx import Document
from django.db.models import Count
from io import BytesIO
from datetime import datetime

def download_word_report(request):
    doc = Document()
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')

    # Форматирование дат
    if start_date:
        start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
    if end_date:
        end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

    # Определение периода и заголовка документа
    if start_date and end_date:
        doc.add_heading(f'Отчет за период: {start_date} - {end_date}', 0)
        receptions = Reception.objects.filter(date__range=(start_date, end_date), comed=True, is_confirmed=True)
    else:
        current_year = datetime.now().year
        doc.add_heading(f'Годовой отчет за {current_year}', 0)
        receptions = Reception.objects.filter(date__year=current_year, comed=True, is_confirmed=True)

    total_patients = receptions.count()
    doc.add_paragraph(f'Общее количество записей: {total_patients}')

    services_stats = receptions.values('services__name').annotate(count=Count('services')).order_by('-count')
    gym_stats = receptions.filter(trainer__isnull=False).values('trainer__name').annotate(count=Count('trainer')).order_by('-count')

    doc.add_heading('Статистика по услугам:', level=1)
    for service in services_stats:
        doc.add_paragraph(f"{service['services__name']}: {service['count']}")

    doc.add_heading('Статистика по тренажерам:', level=1)
    for gym in gym_stats:
        doc.add_paragraph(f"{gym['trainer__name']}: {gym['count']}")

    # Сохраняем документ в памяти
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Отправляем файл пользователю
    response = HttpResponse(file_stream.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="annual_report.docx"'
    return response


@staff_member_required
def unconfirmed_receptions_list(request):
    receptions = Reception.objects.filter(is_confirmed=False).order_by('-date')
    return render(request, 'main/receptions/unconfirmed_receptions_list.html', {'receptions': receptions})

@staff_member_required
def add_reception(request):
    if request.method == 'POST':
        form = ReceptionForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('main/receptions/unconfirmed_receptions_list')
    else:
        form = ReceptionForm()
    return render(request, 'main/receptions/add_reception.html', {'form': form})

from .forms import ReceptionForm
@staff_member_required
def add_reception(request):
    if request.method == 'POST':
        form = ReceptionEditForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('main/receptions/unconfirmed_receptions_list')
    else:
        form = ReceptionEditForm()
    return render(request, 'main/receptions/add_reception.html', {'form': form})

@staff_member_required
def delete_reception(request, pk):
    reception = get_object_or_404(Reception, pk=pk)
    if request.method == 'POST':
        reception.delete()
        return redirect('main:unconfirmed_receptions_list')

    return render(request, 'main/receptions/delete_reception.html', {'reception': reception})

@staff_member_required
def reception_detail(request, pk):
    reception = get_object_or_404(Reception, pk=pk)

    if request.method == 'POST':
        form = ReceptionEditForm(request.POST, instance=reception)
        if form.is_valid():
            form.save()
            return redirect('main:unconfirmed_receptions_list')  # Или другой URL для редиректа
    else:
        form = ReceptionEditForm(instance=reception)

    return render(request, 'main/receptions/reception_detail.html', {
        'form': form,
        'reception': reception
    })
@staff_member_required
def view_reception_detail(request, pk):
    reception = get_object_or_404(Reception, pk=pk)

    return render(request, 'main/receptions/reception_detail_for_doctor.html', {
        'reception': reception
    })

class MedicalEquipmentListView(ListView):
    model = MedicalEquipment
    template_name = 'main/MedicalEquipment/medical_equipment_list.html'  # Обновите путь к вашему шаблону

@staff_member_required
def equipment_detail(request, pk):
    equipment = get_object_or_404(MedicalEquipment, pk=pk)
    return render(request, 'main/MedicalEquipment/equipment_detail.html', {'equipment': equipment})  # Обновите путь к вашему шаблону
class MakeAppointmentView(APIView):
    def post(self, request, *args, **kwargs):
        print("Полученные данные на сервер:", request.data)
        serializer = ReceptionSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            print("Запись на прием успешно создана.")
            return Response({'success': True, 'message': 'Запись на прием успешно создана.'}, status=status.HTTP_201_CREATED)
        else:
            print("Ошибки сериализации:", serializer.errors)
            return Response({'success': False, 'errors': serializer.errors}, status=status.HTTP_400_BAD_REQUEST)


from rest_framework.views import APIView
from rest_framework.response import Response
from .models import Service
from .serializers import ServiceSerializer

class ServiceListViewAPI(APIView):
    def get(self, request, *args, **kwargs):
        services = Service.objects.all()
        serializer = ServiceSerializer(services, many=True, context={'request': request})
        return Response(serializer.data)

class CheckUserView(APIView):
    def post(self, request, *args, **kwargs):
        email = request.data.get('email')
        try:
            user = UserProfile.objects.get(email=email)
            user_data = {
                'full_name': user.full_name,
                'id': user.id,
                # Другие данные, если необходимо
            }
            return Response({
                'is_registered': True,
                'user_info': user_data
            }, status=status.HTTP_200_OK)
        except UserProfile.DoesNotExist:
            # Пользователь не найден
            return Response({'is_registered': False}, status=status.HTTP_200_OK)

class DoctorsListView(APIView):
    def get(self, request):
        doctors = Doctor.objects.all()
        serializer = DoctorSerializer(doctors, many=True)
        return Response(serializer.data)

class ServicesListView(APIView):
    def get(self, request):
        services = Service.objects.all()
        serializer = ServiceSerializer(services, many=True)
        return Response(serializer.data)

class MasseursListView(APIView):
    def get(self, request):
        masseurs = Masseur.objects.all()
        serializer = MasseurSerializer(masseurs, many=True)
        return Response(serializer.data)

class TrainingEquipmentsListView(APIView):
    def get(self, request):
        equipments = TrainingEquipment.objects.all()
        serializer = TrainingEquipmentSerializer(equipments, many=True)
        return Response(serializer.data)